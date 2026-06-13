#!/usr/bin/env python3
"""
将论文正文 Markdown 转换为格式化 Word 文档。
保留模板的封面、原创声明、目录，将正文内容填入对应位置。

用法：
    python md_to_docx.py <input.md> <template.docx> <output.docx>
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("需要安装 python-docx：pip install python-docx")
    sys.exit(1)


# ── 样式映射 ─────────────────────────────────────────────────────────────────
STYLE_H1     = "Heading 1"
STYLE_H2     = "Heading 2"
STYLE_H3     = "Heading 3"
STYLE_BODY   = "Normal Indent"
STYLE_BIBLIO = "Bibliography"
STYLE_KW     = "Normal"


# ── 工具函数 ─────────────────────────────────────────────────────────────────

def para_text(para):
    return "".join(t.text or "" for t in para._element.iter(qn("w:t")))


def replace_text_in_body(doc_body, old_text, new_text):
    """在文档 body（含 textbox）的所有 w:t 元素中替换文本。"""
    for t_elem in doc_body.iter(qn("w:t")):
        if t_elem.text and old_text in t_elem.text:
            t_elem.text = t_elem.text.replace(old_text, new_text)


def strip_heading_number(text, level):
    """剥去 MD 标题中的自动编号前缀，避免与 Word 样式的多级列表重复。"""
    if level == 1:
        # 第一章 / 第二章 … 或阿拉伯数字
        text = re.sub(r"^第[一二三四五六七八九十百零\d]+章\s*", "", text)
    elif level == 2:
        # 1.1 / 2.3 …
        text = re.sub(r"^\d+\.\d+\s+", "", text)
    elif level == 3:
        # 1.1.1 / 2.3.1 …
        text = re.sub(r"^\d+\.\d+\.\d+\s+", "", text)
    return text.strip()


def find_para(doc, contains_text, style_name=None):
    """找到第一个符合文本+样式条件的段落对象。"""
    for p in doc.paragraphs:
        if contains_text in para_text(p):
            if style_name is None or p.style.name == style_name:
                return p
    return None


def find_para_after(doc, anchor_para, contains_text, style_name=None):
    """在 anchor_para 之后找符合条件的段落。"""
    found = False
    for p in doc.paragraphs:
        if p._element is anchor_para._element:
            found = True
            continue
        if found and contains_text in para_text(p):
            if style_name is None or p.style.name == style_name:
                return p
    return None


def paras_between(doc, start_para, end_para):
    """收集 start_para（不含）到 end_para（不含）之间的段落。"""
    result = []
    collecting = False
    for p in doc.paragraphs:
        if p._element is start_para._element:
            collecting = True
            continue
        if end_para is not None and p._element is end_para._element:
            break
        if collecting:
            result.append(p)
    return result


def delete_paragraph(para):
    """从文档中移除一个段落元素。"""
    para._element.getparent().remove(para._element)


def set_page_break_before(para):
    """为段落设置 Word「段前分页」属性。"""
    from docx.oxml import OxmlElement
    pPr = para._element.get_or_add_pPr()
    if pPr.find(qn("w:pageBreakBefore")) is None:
        pbr = OxmlElement("w:pageBreakBefore")
        pPr.append(pbr)


def insert_para_before(anchor_elem, doc, text, style_name):
    """在 anchor_elem 之前插入一个新段落（保证样式正确）。"""
    p = doc.add_paragraph("", style=style_name)
    _fill_para_text(p, text)
    doc.element.body.remove(p._element)
    anchor_elem.addprevious(p._element)
    return p


def insert_para_after(anchor_elem, doc, text, style_name):
    """在 anchor_elem 之后插入一个新段落。"""
    p = doc.add_paragraph("", style=style_name)
    _fill_para_text(p, text)
    doc.element.body.remove(p._element)
    anchor_elem.addnext(p._element)
    return p


def unescape_md(text):
    """移除 Markdown 转义反斜杠，如 \\< → <，\\= → =。"""
    return re.sub(r'\\([\\`*_{}\[\]()#+\-.!<>=])', r'\1', text)


def _fill_para_text(para, text):
    """填充段落文本，支持 **加粗** 行内标记。"""
    text = unescape_md(text)
    for r in list(para._element.findall(qn("w:r"))):
        para._element.remove(r)
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        if i % 2 == 1:
            run.bold = True


# ── Markdown 解析 ─────────────────────────────────────────────────────────────

def parse_md(md_text):
    """将 Markdown 正文解析为各节内容列表：[(text, style), ...]"""
    lines = md_text.splitlines()

    sections = {
        "title":       "",          # 论文中文标题
        "title_en":    "",          # 论文英文标题（# EN: ... 行）
        "abstract_cn": [],
        "abstract_en": [],
        "body":        [],
        "references":  [],
    }

    mode = None
    table_buf = []

    def flush_table():
        """将收集的表格行转为占位段落。"""
        if table_buf:
            sections[mode].append(("[⚠️ 此处为表格，请在 Word 中手动插入]", STYLE_BODY))
            table_buf.clear()

    for line in lines:
        stripped = line.strip()

        # 跳过元信息行、分割线
        if (not stripped
                or stripped.startswith("> ")
                or re.match(r'^[-*_]{3,}$', stripped.replace(" ", ""))):
            if table_buf and mode:
                flush_table()
            continue

        # 论文大标题（# ，不是 ##）→ 提取标题文本
        if re.match(r"^# EN:", stripped):
            sections["title_en"] = stripped[5:].strip()
            continue
        if re.match(r"^# [^#]", stripped):
            sections["title"] = stripped[2:].strip()
            continue

        # ── 节头检测 ──────────────────────────────────────────────────────────
        if stripped == "## 摘要" or stripped.startswith("## 摘要"):
            mode = "abstract_cn"
            continue

        if stripped == "**Abstract**":
            if table_buf:
                flush_table()
            mode = "abstract_en"
            continue

        # ## 第X章 或纯数字章节头
        if re.match(r"^## 第[一二三四五六七八九十百\d]+章", stripped):
            if table_buf:
                flush_table()
            mode = "body"
            # 这行本身是 H1，继续向下处理

        if re.match(r"^## 参考文献", stripped):
            if table_buf:
                flush_table()
            mode = "references"
            continue

        # ── 表格行 ────────────────────────────────────────────────────────────
        if stripped.startswith("|") and mode in ("body",):
            table_buf.append(stripped)
            continue
        else:
            if table_buf:
                flush_table()

        # ── 按模式分发 ────────────────────────────────────────────────────────
        if mode == "abstract_cn":
            if stripped.startswith("**关键词") or stripped.startswith("**Key"):
                sections["abstract_cn"].append((stripped, STYLE_KW))
            elif stripped:
                sections["abstract_cn"].append((stripped, STYLE_BODY))

        elif mode == "abstract_en":
            if stripped.startswith("**Keywords") or stripped.startswith("Keywords"):
                sections["abstract_en"].append((stripped, STYLE_KW))
            elif stripped:
                sections["abstract_en"].append((stripped, STYLE_BODY))

        elif mode == "body":
            if re.match(r"^## 第[一二三四五六七八九十百\d]+章", stripped):
                title = strip_heading_number(stripped[3:], level=1)
                sections["body"].append((title, STYLE_H1))
            elif stripped.startswith("### "):
                title = strip_heading_number(stripped[4:], level=2)
                sections["body"].append((title, STYLE_H2))
            elif stripped.startswith("#### "):
                title = strip_heading_number(stripped[5:], level=3)
                sections["body"].append((title, STYLE_H3))
            elif stripped.startswith("##### "):
                title = strip_heading_number(stripped[6:], level=3)
                sections["body"].append((title, STYLE_H3))
            elif stripped.startswith(("- ", "* ", "· ")):
                sections["body"].append(("• " + stripped[2:], STYLE_BODY))
            elif re.match(r"^\d+\.\s", stripped):
                sections["body"].append((stripped, STYLE_BODY))
            elif stripped:
                sections["body"].append((stripped, STYLE_BODY))

        elif mode == "references":
            if stripped:
                sections["references"].append((stripped, STYLE_BIBLIO))

    if table_buf:
        flush_table()

    return sections


# ── 主转换逻辑 ────────────────────────────────────────────────────────────────

def convert(input_md, template_docx, output_docx):
    input_path   = Path(input_md)
    template_path = Path(template_docx)
    output_path  = Path(output_docx)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if not input_path.exists():
        print(f"错误：找不到输入文件 {input_path}"); sys.exit(1)
    if not template_path.exists():
        print(f"错误：找不到模板文件 {template_path}"); sys.exit(1)

    print(f"读取模板：{template_path}")
    doc = Document(str(template_path))

    print(f"解析正文：{input_path}")
    md_text  = input_path.read_text(encoding="utf-8")
    sections = parse_md(md_text)

    # ── 1. 定位模板关键段落 ────────────────────────────────────────────────────
    abstract_cn_h = find_para(doc, "摘",        style_name="标题 1*")
    abstract_en_h = find_para(doc, "Abstract",  style_name="标题 1*")
    toc_h         = find_para(doc, "目",         style_name="标题 1*")
    refs_h        = find_para(doc, "参考文献",   style_name="标题 1*")

    # 找目录之后的第一个正文 Heading 1（占位内容起点）
    first_body_h = None
    found_toc = False
    for p in doc.paragraphs:
        if p._element is toc_h._element:
            found_toc = True
            continue
        if found_toc and p.style.name == "Heading 1":
            first_body_h = p
            break

    # 找模板保留章节：「图表、公式格式」（正文末尾备用章，不删除）
    fig_table_h = None
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and "图表" in para_text(p):
            fig_table_h = p
            break

    # 找附录/致谢（参考文献之后的后置内容起点）
    back_matter_h = None
    found_refs = False
    for p in doc.paragraphs:
        if refs_h and p._element is refs_h._element:
            found_refs = True
            continue
        if found_refs and p.style.name in ("附录 1", "标题 1*"):
            back_matter_h = p
            break

    fbh  = para_text(first_body_h)  if first_body_h  else "N/A"
    fth  = para_text(fig_table_h)   if fig_table_h   else "未找到"
    rh   = para_text(refs_h)        if refs_h        else "N/A"
    bmh  = para_text(back_matter_h) if back_matter_h else "末尾"
    print(f"  摘要 CN  : [{para_text(abstract_cn_h)}]")
    print(f"  摘要 EN  : [{para_text(abstract_en_h)}]")
    print(f"  目录     : [{para_text(toc_h)}]")
    print(f"  正文起点 : [{fbh}]")
    print(f"  备用章节 : [{fth}]")
    print(f"  参考文献 : [{rh}]")
    print(f"  后置内容 : [{bmh}]")

    # ── 2. 各节标题加「段前分页」 ────────────────────────────────────────────
    for h in [abstract_cn_h, abstract_en_h, toc_h, refs_h]:
        if h:
            set_page_break_before(h)

    # ── 4. 替换封面标题占位文本 ───────────────────────────────────────────────
    COVER_TITLE_PLACEHOLDER_CN = "上海交通大学学位论文格式模板"
    COVER_TITLE_PLACEHOLDER_EN = "Dissertation Template for Master Degree of Engineering in Shanghai Jiao Tong University"

    if sections["title"]:
        replace_text_in_body(doc.element.body, COVER_TITLE_PLACEHOLDER_CN, sections["title"])
        print(f"  中文封面替换：{sections['title']!r}")

    if sections["title_en"]:
        replace_text_in_body(doc.element.body, COVER_TITLE_PLACEHOLDER_EN, sections["title_en"])
        print(f"  英文封面替换：{sections['title_en']!r}")
    else:
        print("  英文封面：未找到 '# EN: ...' 行，保留模板占位文本")

    # ── 5. 删除占位内容（先收集再删除，避免迭代中修改） ──────────────────────
    to_delete = []

    # 摘要 CN 占位（摘要 heading 与 Abstract heading 之间）
    to_delete += paras_between(doc, abstract_cn_h, abstract_en_h)

    # 摘要 EN 占位（Abstract heading 与 目录 heading 之间）
    to_delete += paras_between(doc, abstract_en_h, toc_h)

    # 正文占位（first_body_h 到「图表」备用章 或 参考文献 heading 之间）
    # 「图表、公式格式」章之后的内容作为备用保留，不删除
    body_end = fig_table_h if fig_table_h else refs_h
    if first_body_h and body_end:
        to_delete.append(first_body_h)
        to_delete += paras_between(doc, first_body_h, body_end)

    # 参考文献占位条目（参考文献 heading 与 后置内容 之间）
    if refs_h:
        to_delete += paras_between(doc, refs_h, back_matter_h)

    for p in to_delete:
        delete_paragraph(p)

    print(f"  已删除占位段落：{len(to_delete)} 段")

    # ── 6. 插入新内容 ──────────────────────────────────────────────────────────
    PLACEHOLDER_NOTE = "【备用章节：以下为模板示例，正式提交前请删除或替换】"

    # 摘要 CN → 插入到 Abstract heading 之前
    for text, style in sections["abstract_cn"]:
        insert_para_before(abstract_en_h._element, doc, text, style)

    # 摘要 EN → 插入到 目录 heading 之前
    for text, style in sections["abstract_en"]:
        insert_para_before(toc_h._element, doc, text, style)

    # 正文章节 → 插入到「图表」备用章（或参考文献）之前
    # Heading 1 段落加「段前分页」
    body_anchor = fig_table_h if fig_table_h else refs_h
    if body_anchor:
        for text, style in sections["body"]:
            p = insert_para_before(body_anchor._element, doc, text, style)
            if style == STYLE_H1:
                set_page_break_before(p)

    # 备用章节标注：「图表、公式格式」
    if fig_table_h:
        insert_para_after(fig_table_h._element, doc, PLACEHOLDER_NOTE, "Normal Indent")
        set_page_break_before(fig_table_h)   # 备用章也加分页，视觉上分隔清楚

    # 参考文献条目 → 插入到 后置内容 之前（或末尾）
    if sections["references"]:
        anchor = back_matter_h._element if back_matter_h else None
        for text, style in sections["references"]:
            if anchor is not None:
                insert_para_before(anchor, doc, text, style)
            else:
                p = doc.add_paragraph("", style=style)
                _fill_para_text(p, text)

    # 备用章节标注：「实验环境」
    if back_matter_h:
        insert_para_after(back_matter_h._element, doc, PLACEHOLDER_NOTE, "Normal Indent")

    print(f"保存输出：{output_path}")
    doc.save(str(output_path))
    print(f"✅ 完成：{output_path}")
    print(f"   摘要CN {len(sections['abstract_cn'])}段 | "
          f"摘要EN {len(sections['abstract_en'])}段 | "
          f"正文 {len(sections['body'])}段 | "
          f"参考文献 {len(sections['references'])}条")


if __name__ == "__main__":
    if len(sys.argv) < 4:
        print(__doc__)
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2], sys.argv[3])
